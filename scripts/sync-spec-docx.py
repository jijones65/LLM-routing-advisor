#!/usr/bin/env python3
"""Render SPECIFICATION.md as the repository's authoritative Word specification."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "SPECIFICATION.md"
OUTPUT = ROOT / "SPEC_AI routing advisor.docx"
USABLE_DXA = 9360
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    if hasattr(element, "_tc"):
        props = element._tc.get_or_add_tcPr()
    elif element.tag.endswith("}tc"):
        props = element.get_or_add_tcPr()
    else:
        props = element.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    text_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    rpr.extend([color, size])
    text_run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    text_run.append(text)
    field.append(text_run)
    paragraph._p.append(field)


def remove_paragraph_border(style):
    p_pr = style._element.get_or_add_pPr()
    border = p_pr.find(qn("w:pBdr"))
    if border is not None:
        p_pr.remove(border)


def next_numbering_id(numbering, tag, attribute):
    values = [int(node.get(qn(attribute), "0")) for node in numbering.findall(qn(tag))]
    return max(values, default=0) + 1


def add_list_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level_justify = OxmlElement("w:lvlJc")
    level_justify.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    level.extend([start, number_format, level_text, suffix, level_justify, p_pr])
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def add_list_number(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    num_id = next_numbering_id(numbering, "w:num", "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    restart = OxmlElement("w:startOverride")
    restart.set(qn("w:val"), "1")
    override.append(restart)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_list_number(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.append(num_pr)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(27)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    remove_paragraph_border(title)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
        ("Heading 4", 11, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    metadata = styles.add_style("Spec Metadata", WD_STYLE_TYPE.PARAGRAPH)
    metadata.base_style = normal
    metadata.font.name = "Calibri"
    metadata.font.size = Pt(10)
    metadata.font.color.rgb = RGBColor.from_string(MUTED)
    metadata.paragraph_format.space_after = Pt(3)

    callout = styles.add_style("Spec Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout.base_style = normal
    callout.font.name = "Calibri"
    callout.font.size = Pt(11)
    callout.font.italic = True
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.left_indent = Inches(0.2)
    callout.paragraph_format.right_indent = Inches(0.15)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)

    code = styles.add_style("Spec Code", WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = normal
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("LLM Application Routing Advisor")
    set_font(left, size=9, color=MUTED, bold=True)
    right = header.add_run("  ·  Product and Implementation Specification")
    set_font(right, size=9, color=MUTED)
    header.paragraph_format.space_after = Pt(0)
    add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = "LLM Application Routing Advisor - Product and Implementation Specification"
    doc.core_properties.subject = "Living product and implementation specification"
    doc.core_properties.keywords = "LLM, model routing, application design, validation, specification"


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|(?<!\*)\*[^*]+\*(?!\*))")


def add_hyperlink(paragraph, label, url):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = label
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, content, default_size=11):
    cursor = 0
    for match in INLINE.finditer(content):
        if match.start() > cursor:
            run = paragraph.add_run(content[cursor : match.start()])
            set_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=default_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=max(8.5, default_size - 1), color=DARK_BLUE)
        elif token.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=default_size, italic=True)
        cursor = match.end()
    if cursor < len(content):
        run = paragraph.add_run(content[cursor:])
        set_font(run, size=default_size)


def table_cells(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line):
    return bool(line.strip().startswith("|") and all(re.fullmatch(r":?-{3,}:?", cell) for cell in table_cells(line)))


def table_widths(rows):
    columns = len(rows[0])
    weights = []
    for index in range(columns):
        longest = max(len(row[index]) if index < len(row) else 0 for row in rows)
        weights.append(max(8, min(longest, 42)))
    minimum = 720 if columns <= 6 else 540
    remaining = USABLE_DXA - minimum * columns
    total_weight = sum(weights)
    widths = [minimum + int(remaining * weight / total_weight) for weight in weights]
    widths[-1] += USABLE_DXA - sum(widths)
    return widths


def add_table(doc, rows):
    columns = len(rows[0])
    normalised = [row[:columns] + [""] * max(0, columns - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalised), cols=columns)
    table.style = "Table Grid"
    widths = table_widths(normalised)
    set_table_geometry(table, widths)
    small = 8.5 if columns >= 6 else 9.5
    for row_index, row in enumerate(normalised):
        for col_index, content in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, content, default_size=small)
            if row_index == 0:
                shade(cell, TABLE_FILL)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    index = 0
    in_code = False
    code_lines = []
    title_seen = False
    bullet_num_id = None
    number_num_id = None
    list_abstracts = {
        "bullet": add_list_definition(doc, "bullet"),
        "number": add_list_definition(doc, "number"),
    }

    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                for code_line in code_lines or [""]:
                    paragraph = doc.add_paragraph(style="Spec Code")
                    paragraph.paragraph_format.keep_together = True
                    shade(paragraph._p, CALLOUT_FILL)
                    run = paragraph.add_run(code_line or " ")
                    set_font(run, name="Consolas", size=9, color="222222")
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if not bullet and not numbered:
            bullet_num_id = None
            number_num_id = None
        if not line.strip():
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows = [table_cells(line)]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(table_cells(lines[index]))
                index += 1
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            content = heading.group(2).strip()
            if level == 1 and not title_seen:
                paragraph = doc.add_paragraph(style="Title")
                add_inline(paragraph, content, default_size=27)
                subtitle = doc.add_paragraph("Living specification for refinement, maintenance and future rebuilds", style="Subtitle")
                subtitle.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
                title_seen = True
            else:
                paragraph = doc.add_paragraph(style=f"Heading {min(max(level - 1, 1), 4)}")
                add_inline(paragraph, content, default_size={2: 16, 3: 13, 4: 12}.get(level, 11))
            index += 1
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph(style="Spec Callout")
            shade(paragraph._p, CALLOUT_FILL)
            add_inline(paragraph, line[2:].strip())
            index += 1
            continue
        if bullet:
            if bullet_num_id is None:
                bullet_num_id = add_list_number(doc, list_abstracts["bullet"])
            paragraph = doc.add_paragraph(style="Normal")
            paragraph.paragraph_format.space_after = Pt(4)
            apply_list_number(paragraph, bullet_num_id)
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        if numbered:
            if number_num_id is None:
                number_num_id = add_list_number(doc, list_abstracts["number"])
            bullet_num_id = None
            paragraph = doc.add_paragraph(style="Normal")
            paragraph.paragraph_format.space_after = Pt(4)
            apply_list_number(paragraph, number_num_id)
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        style = "Spec Metadata" if title_seen and index < 8 and line.startswith("**") else "Normal"
        paragraph = doc.add_paragraph(style=style)
        add_inline(paragraph, line.strip(), default_size=10 if style == "Spec Metadata" else 11)
        index += 1

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT} from {SOURCE} ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")


if __name__ == "__main__":
    build()
